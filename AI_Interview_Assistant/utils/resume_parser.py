# utils/resume_parser.py
"""
Resume Parser Utility.
Handles raw text extraction from PDF/DOCX and parses basic sections (Contact, Education, Experience).
"""

import re
import pdfplumber
import docx

class ResumeParser:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extracts all raw text from a PDF file using pdfplumber.
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extracts all raw text from a DOCX file using python-docx.
        """
        text = []
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text.append(para.text)
            # Process tables as well
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
        except Exception as e:
            print(f"Error reading DOCX: {e}")
        return "\n".join(text)

    @staticmethod
    def extract_contact_info(text: str) -> dict:
        """
        Extracts email, phone, and name from the resume text.
        """
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        
        email_match = re.search(email_pattern, text)
        phone_match = re.search(phone_pattern, text)
        
        email = email_match.group(0) if email_match else "Not Found"
        phone = phone_match.group(0) if phone_match else "Not Found"
        
        # Name heuristic: typically in the first few lines of the text,
        # let's clean the first non-empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        name = "Candidate"
        if lines:
            # Pick first line, if it doesn't look like an email/phone/label
            first_line = lines[0]
            if len(first_line) < 40 and not re.search(email_pattern, first_line) and not re.search(phone_pattern, first_line):
                name = first_line
        
        return {
            "name": name,
            "email": email,
            "phone": phone
        }

    @staticmethod
    def extract_education(text: str) -> list:
        """
        Searches for lines containing common degrees or education keywords.
        """
        education_keywords = [
            "b.tech", "btech", "b.e.", "be", "m.tech", "mtech", "b.sc", "bsc", 
            "m.sc", "msc", "mca", "bca", "mba", "ph.d", "phd", "bachelor", "master", 
            "university", "college", "institute", "school", "education", "academic"
        ]
        
        education_records = []
        lines = text.split('\n')
        
        # Scan line-by-line for academic mentions
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(kw in line_lower for kw in education_keywords):
                cleaned = line.strip()
                if len(cleaned) > 5 and len(cleaned) < 150:
                    # Capture surrounding context (like next line which might have university name or date)
                    context = cleaned
                    if i + 1 < len(lines):
                        next_line = lines[i+1].strip()
                        if next_line and len(next_line) < 100 and not any(kw in next_line.lower() for kw in ["experience", "work", "skills"]):
                            context += " - " + next_line
                    if context not in education_records:
                        education_records.append(context)
                        
        return education_records[:4] # Limit to top 4 matches to prevent bloating

    @staticmethod
    def extract_experience(text: str) -> list:
        """
        Heuristic-based extraction of job experience or project descriptions.
        """
        experience_headers = [
            "experience", "work experience", "professional experience", "employment history",
            "projects", "academic projects", "employment", "internship", "intern"
        ]
        
        experience_records = []
        lines = text.split('\n')
        recording = False
        current_record = []
        
        for line in lines:
            line_strip = line.strip()
            if not line_strip:
                continue
                
            line_lower = line_strip.lower()
            # Detect section header
            if any(line_lower.startswith(header) or line_lower == header for header in experience_headers):
                recording = True
                if current_record:
                    experience_records.append("\n".join(current_record))
                    current_record = []
                continue
                
            # If recording and we hit another major section header like Skills or Education, stop recording
            if recording and any(line_lower.startswith(sect) for sect in ["skills", "education", "certifications", "interests", "languages"]):
                recording = False
                if current_record:
                    experience_records.append("\n".join(current_record))
                    current_record = []
                
            if recording:
                if len(current_record) < 15: # Grab up to 15 lines per section
                    current_record.append(line_strip)
                else:
                    recording = False
                    experience_records.append("\n".join(current_record))
                    current_record = []
                    
        if current_record:
            experience_records.append("\n".join(current_record))
            
        # Clean up records and keep non-empty, reasonably long ones
        cleaned_records = []
        for rec in experience_records:
            cleaned_rec = rec.strip()
            if len(cleaned_rec) > 30:
                cleaned_records.append(cleaned_rec)
                
        return cleaned_records[:3] # Limit to top 3 entries
